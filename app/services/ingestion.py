from pypdf import PdfReader
from docx import Document as DocxDocument
from io import BytesIO
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session
from app.models.models import Document, DocumentChunk
from app.services.embeddings import embed_texts

# Chunking config — these numbers matter
CHUNK_SIZE = 500      # characters per chunk
CHUNK_OVERLAP = 50    # overlap between chunks to preserve context

splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", ". ", " ", ""]
)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text.strip())
    return "\n\n".join(pages)

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)

def extract_text_from_txt(file_bytes: bytes) -> str:
    # Try UTF-8 first, fall back to latin-1
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")

def ingest_document(
    filename: str,
    file_bytes: bytes,
    tenant_id: str,
    db: Session
) -> dict:
    # Step 1: Extract text based on file extension
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
        if not raw_text.strip():
            raise ValueError("Could not extract text from PDF")
    elif ext == "docx":
        raw_text = extract_text_from_docx(file_bytes)
        if not raw_text.strip():
            raise ValueError("Could not extract text from DOCX")
    elif ext == "txt":
        raw_text = extract_text_from_txt(file_bytes)
        if not raw_text.strip():
            raise ValueError("The text file appears to be empty")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    # Step 2: Chunk
    chunks = splitter.split_text(raw_text)
    if not chunks:
        raise ValueError("Document produced no chunks after splitting")

    # Step 3: Embed in small batches to avoid OOM on free tier
    BATCH_SIZE = 16
    embeddings = []
    for i in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[i:i + BATCH_SIZE]
        embeddings.extend(embed_texts(batch))

    # Step 4: Save document record
    document = Document(
        filename=filename,
        tenant_id=tenant_id
    )
    db.add(document)
    db.flush()  # get document.id before committing

    # Step 5: Save all chunks with tenant_id scoped directly
    chunk_records = [
        DocumentChunk(
            content=chunk,
            chunk_index=i,
            embedding=embedding,
            document_id=document.id,
            tenant_id=tenant_id  # ← directly on every chunk
        )
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings))
    ]
    db.bulk_save_objects(chunk_records)
    db.commit()

    return {
        "document_id": document.id,
        "filename": filename,
        "chunks_created": len(chunks),
        "tenant_id": tenant_id
    }